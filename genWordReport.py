# -*- coding: utf8 -*-
# coding=utf-8 

import sys
import win32com.client
from docx import Document

reload(sys)
sys.setdefaultencoding('utf-8')

document = Document()
'''
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_heading('Heading, level 1', level=1)
document.add_heading('Heading, level 1', level=1)
document.add_heading('Heading, level 1', level=1)
#document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
	'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
	'first item in ordered list', style='ListNumber'
)

#document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

items = [[1, 101, 'Spam'], [2, 42, 'Eggs'], [3, 631, 'Spam, spam']]

for item in items:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item[0])
	row_cells[1].text = str(item[1])
	row_cells[2].text = item[2]

document.add_page_break()
'''

wordFilePath = "../RGtool/report/hello.docx"
excelFilePath = "../RGtool/resource/test.xlsx"

excelapp = win32com.client.Dispatch("Excel.Application")
excelapp.Visible = 0
excelxls = excelapp.Workbooks.Open(excelFilePath)

ws = excelxls.Worksheets("titles")
used = ws.UsedRange
nrows = used.Row + used.Rows.Count
ncols = used.Column + used.Columns.Count

'''
for i in range(1, nrows):
	for j in range(1, ncols):
		print ws.Cells(i, j)
'''
for i in range(2, nrows):
	#document.add_heading(str(ws.Cells(i, 1)), level=int(ws.Cells(i, 2)))
	level = int(ws.Cells(i, 2))
	if level == 1 :
		document.add_paragraph(str(ws.Cells(i, 1)), style="List Number")
	else :
		document.add_paragraph(str(ws.Cells(i, 1)), style="List Number " + str(level))

ws = excelxls.Worksheets("vuls")
used = ws.UsedRange
nrows = used.Row + used.Rows.Count
ncols = used.Column + used.Columns.Count

for i in range(2, nrows):
	for j in range(2, ncols):
		line = str(ws.Cells(1, j)) + " : " + unicode(ws.Cells(i, j))
		document.add_paragraph(line, style='List Bullet')

document.save(wordFilePath)
excelapp.Quit() # Close the Word Application
